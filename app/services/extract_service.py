# ══════════════════════════════════════════════════════════
# app/services/extract_service.py
# Extracción de texto de PDF, Word e Imágenes para usar como
# "documento de referencia" en el Redactor.
# ══════════════════════════════════════════════════════════

import io
import os
import base64
from typing import Optional, Tuple

# ─── Detección de librerías opcionales ───
try:
    import pdfplumber  # type: ignore
    _HAS_PDFPLUMBER = True
except Exception:
    _HAS_PDFPLUMBER = False

try:
    import docx  # python-docx
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False

try:
    from openai import OpenAI  # type: ignore
    _HAS_OPENAI = True
except Exception:
    _HAS_OPENAI = False


# ─── Tipos soportados ───
EXT_PDF = {".pdf"}
EXT_DOCX = {".docx", ".doc"}
EXT_IMG = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
EXT_TXT = {".txt", ".md"}

EXTS_SOPORTADAS = EXT_PDF | EXT_DOCX | EXT_IMG | EXT_TXT

MAX_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_CHARS_OUT = 8000  # truncamos lo que devolvemos


def _ext(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def soportado(filename: str) -> bool:
    return _ext(filename) in EXTS_SOPORTADAS


def extraer_texto(
    filename: str,
    contenido: bytes,
    api_key: Optional[str] = None,
) -> Tuple[str, Optional[str]]:
    """
    Extrae texto del archivo. Devuelve (texto, error).
    Si no hay error, error=None. Si hay error, texto="" y error con motivo.
    """
    if not filename or not contenido:
        return "", "Archivo vacío"

    if len(contenido) > MAX_BYTES:
        return "", f"Archivo demasiado grande (máximo {MAX_BYTES // (1024*1024)} MB)"

    ext = _ext(filename)
    if ext not in EXTS_SOPORTADAS:
        return "", f"Tipo no soportado: {ext or 'desconocido'}"

    try:
        if ext in EXT_TXT:
            return _truncar(contenido.decode("utf-8", errors="ignore")), None

        if ext in EXT_PDF:
            return _extraer_pdf(contenido)

        if ext in EXT_DOCX:
            return _extraer_docx(contenido)

        if ext in EXT_IMG:
            return _extraer_imagen(contenido, ext, api_key)
    except Exception as e:
        return "", f"Error al extraer texto: {e}"

    return "", "Tipo no soportado"


def _truncar(texto: str) -> str:
    texto = (texto or "").strip()
    if len(texto) > MAX_CHARS_OUT:
        return texto[:MAX_CHARS_OUT] + "\n[...truncado...]"
    return texto


def _extraer_pdf(contenido: bytes) -> Tuple[str, Optional[str]]:
    if not _HAS_PDFPLUMBER:
        return "", "pdfplumber no instalado en el servidor"
    partes = []
    with pdfplumber.open(io.BytesIO(contenido)) as pdf:
        for page in pdf.pages[:30]:  # máximo 30 páginas
            t = page.extract_text() or ""
            if t.strip():
                partes.append(t)
    texto = "\n\n".join(partes).strip()
    if not texto:
        return "", "El PDF no contiene texto extraíble (¿es escaneado?)"
    return _truncar(texto), None


def _extraer_docx(contenido: bytes) -> Tuple[str, Optional[str]]:
    if not _HAS_DOCX:
        return "", "python-docx no instalado en el servidor"
    doc = docx.Document(io.BytesIO(contenido))
    parrafos = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # Tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text and c.text.strip()]
            if celdas:
                parrafos.append(" | ".join(celdas))
    texto = "\n".join(parrafos).strip()
    if not texto:
        return "", "El documento Word está vacío"
    return _truncar(texto), None


def _extraer_imagen(
    contenido: bytes, ext: str, api_key: Optional[str]
) -> Tuple[str, Optional[str]]:
    """
    Usa GPT-4o Vision para describir / leer texto de la imagen.
    Si no hay API key, devolvemos un placeholder informando que se subió
    una imagen pero no se pudo extraer.
    """
    key = api_key or os.environ.get("OPENAI_API_KEY")
    if not _HAS_OPENAI or not key:
        return (
            "[Imagen adjunta. La IA no está conectada — configura OPENAI_API_KEY "
            "en Railway para que GPT-4o Vision lea el contenido de la imagen.]",
            None,
        )

    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
    }
    mime = mime_map.get(ext, "image/png")
    b64 = base64.b64encode(contenido).decode("ascii")
    data_url = f"data:{mime};base64,{b64}"

    try:
        client = OpenAI(api_key=key)
        resp = client.chat.completions.create(
            model="gpt-4o",
            temperature=0.0,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Eres un OCR + descriptor de documentos institucionales peruanos. "
                        "Extrae TODO el texto visible (encabezados, oficios, fechas, números de "
                        "documento, firmas, sellos) y haz un breve resumen al final. "
                        "Responde en texto plano, sin markdown."
                    ),
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "Extrae el texto visible y los datos institucionales de esta imagen. "
                                "Si es un oficio, indica número, fecha, remitente, destinatario y asunto."
                            ),
                        },
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                },
            ],
        )
        texto = (resp.choices[0].message.content or "").strip()
        if not texto:
            return "", "Vision no devolvió texto"
        return _truncar(texto), None
    except Exception as e:
        return "", f"Error en Vision: {e}"
