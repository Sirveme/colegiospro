# ══════════════════════════════════════════════════════════
# app/services/docx_service.py
# Generación de archivo .docx con membrete institucional.
# ══════════════════════════════════════════════════════════

import re
from io import BytesIO
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False


URL_RE = re.compile(r'(https?://[^\s]+)')
_MD_BOLD_RE    = re.compile(r'\*\*(.+?)\*\*')
_MD_OL_RE      = re.compile(r'^\s*\d+\.\s+')
_MD_TBL_SEP_RE = re.compile(r'^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$')
_HTML_HINT_RE  = re.compile(r'<\s*(p|div|table|ul|ol|li|tr|td|th|strong|b|em|i|br|h[1-6])\b', re.I)

try:
    from bs4 import BeautifulSoup, NavigableString
    _HAS_BS4 = True
except Exception:
    _HAS_BS4 = False


def es_contenido_html(texto: str) -> bool:
    if not texto:
        return False
    return bool(_HTML_HINT_RE.search(texto))


def _limpiar_delimitadores(texto: str) -> str:
    """Quita delimitadores ''' o ``` al inicio/fin."""
    texto = (texto or "").strip()
    for delim in ("'''", '"""', "```"):
        if texto.startswith(delim):
            texto = texto[len(delim):]
        if texto.endswith(delim):
            texto = texto[:-len(delim)]
    return texto.strip()


def _agregar_hipervinculo(paragraph, url: str, texto_visible: str = ""):
    """Inserta un hipervínculo azul subrayado en el párrafo."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = texto_visible or url
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _agregar_parrafo_con_links(doc, texto_linea: str):
    """Crea un párrafo y divide el texto entre texto plano y URLs activas."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)

    if not texto_linea:
        return p

    partes = URL_RE.split(texto_linea)
    for parte in partes:
        if not parte:
            continue
        if URL_RE.fullmatch(parte):
            _agregar_hipervinculo(p, parte, parte)
        else:
            p.add_run(parte)
    return p


# Márgenes por tipo de documento — alineados con pdf_service
MARGENES_POR_TIPO = {
    "oficio":              {"izq": 3.5, "der": 2.5},
    "oficio_multiple":     {"izq": 3.5, "der": 2.5},
    "carta":               {"izq": 3.0, "der": 2.5},
    "memorandum":          {"izq": 2.5, "der": 2.5},
    "memorandum_multiple": {"izq": 2.5, "der": 2.5},
    "resolucion":          {"izq": 3.0, "der": 2.5},
    "acta":                {"izq": 2.5, "der": 2.5},
    "circular":            {"izq": 3.0, "der": 2.5},
    "comunicado_general":  {"izq": 2.5, "der": 2.5},
    "orden_pedido":        {"izq": 2.5, "der": 2.5},
}


def docx_disponible() -> bool:
    return _HAS_DOCX


def _parsear_celdas_md(linea: str) -> list:
    s = (linea or '').strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]
    return [c.strip() for c in s.split('|')]


def _agregar_runs_con_negrita(parrafo, texto: str):
    """Divide el texto por **...** y agrega runs alternando bold."""
    if not texto:
        return
    partes = _MD_BOLD_RE.split(texto)
    for j, parte in enumerate(partes):
        if not parte:
            continue
        run = parrafo.add_run(parte)
        run.bold = (j % 2 == 1)


def _emitir_inline_html(parrafo, nodo, hereda_bold=False, hereda_italic=False, hereda_under=False):
    """Recorre un nodo bs4 inline y crea runs en el párrafo respetando estilos."""
    if not _HAS_BS4 or nodo is None:
        return
    for hijo in getattr(nodo, "children", []):
        if isinstance(hijo, NavigableString):
            run = parrafo.add_run(str(hijo))
            if hereda_bold:   run.bold = True
            if hereda_italic: run.italic = True
            if hereda_under:  run.underline = True
            continue
        nm = (hijo.name or "").lower()
        if nm == "br":
            parrafo.add_run().add_break()
            continue
        bold = hereda_bold or nm in ("strong", "b")
        ital = hereda_italic or nm in ("em", "i")
        under = hereda_under or nm == "u"
        if nm == "a":
            href = (hijo.get("href") or "").strip()
            texto_link = hijo.get_text() or href
            if href:
                _agregar_hipervinculo(parrafo, href, texto_link)
                continue
        _emitir_inline_html(parrafo, hijo, bold, ital, under)


def html_a_docx(doc, html: str):
    """Convierte HTML del editor Quill a contenido Word."""
    if not _HAS_DOCX or not _HAS_BS4 or not html:
        return
    soup = BeautifulSoup(html, "html.parser")
    raiz = soup.body or soup

    def render_lista(elem, ordenada: bool):
        estilo = "List Number" if ordenada else "List Bullet"
        for li in elem.find_all("li", recursive=False):
            try:
                p = doc.add_paragraph(style=estilo)
            except KeyError:
                p = doc.add_paragraph()
            _emitir_inline_html(p, li)

    def render_tabla(tbl):
        filas_html = tbl.find_all("tr")
        if not filas_html:
            return
        celdas_por_fila = [tr.find_all(["th", "td"]) for tr in filas_html]
        cols = max((len(c) for c in celdas_por_fila), default=0)
        if cols == 0:
            return
        tabla = doc.add_table(rows=len(filas_html), cols=cols)
        try:
            tabla.style = "Table Grid"
        except KeyError:
            pass
        es_header_row = bool(filas_html[0].find("th"))
        for r, celdas in enumerate(celdas_por_fila):
            for c in range(cols):
                cell = tabla.cell(r, c)
                cell.text = ""
                p = cell.paragraphs[0]
                if c < len(celdas):
                    _emitir_inline_html(p, celdas[c])
                if r == 0 and es_header_row:
                    for run in p.runs:
                        run.bold = True
        doc.add_paragraph()

    for elem in raiz.children:
        if isinstance(elem, NavigableString):
            txt = str(elem).strip()
            if txt:
                doc.add_paragraph(txt)
            continue
        nm = (elem.name or "").lower()
        if nm in ("p", "div"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(14)
            _emitir_inline_html(p, elem)
        elif nm in ("h1", "h2", "h3", "h4", "h5", "h6"):
            p = doc.add_paragraph()
            _emitir_inline_html(p, elem)
            for run in p.runs:
                run.bold = True
        elif nm == "ul":
            render_lista(elem, ordenada=False)
        elif nm == "ol":
            render_lista(elem, ordenada=True)
        elif nm == "table":
            render_tabla(elem)
        elif nm == "br":
            doc.add_paragraph()
        elif nm == "hr":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            _emitir_inline_html(p, elem)


def markdown_a_docx(doc, texto: str):
    """Convierte Markdown del documento (tablas, listas, **negrita**) a Word."""
    if not _HAS_DOCX:
        return
    lineas = (texto or '').split('\n')
    i = 0
    n = len(lineas)
    while i < n:
        linea = lineas[i]
        stripped = linea.strip()

        # Tabla Markdown: cabecera + separador |---|
        if stripped.startswith('|') and i + 1 < n and _MD_TBL_SEP_RE.match(lineas[i+1]):
            cabecera = _parsear_celdas_md(lineas[i])
            i += 2
            filas = [cabecera]
            while i < n and lineas[i].strip().startswith('|'):
                filas.append(_parsear_celdas_md(lineas[i]))
                i += 1
            cols = max(len(f) for f in filas)
            for f in filas:
                while len(f) < cols:
                    f.append('')
            tabla = doc.add_table(rows=len(filas), cols=cols)
            try:
                tabla.style = 'Table Grid'
            except KeyError:
                pass
            for r, fila in enumerate(filas):
                for c, celda in enumerate(fila):
                    cell = tabla.cell(r, c)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    _agregar_runs_con_negrita(p, celda)
                    if r == 0:
                        for run in p.runs:
                            run.bold = True
            doc.add_paragraph()
            continue

        # Lista numerada
        if _MD_OL_RE.match(linea):
            cuerpo = _MD_OL_RE.sub('', linea, count=1)
            try:
                p = doc.add_paragraph(style='List Number')
            except KeyError:
                p = doc.add_paragraph()
            _agregar_runs_con_negrita(p, cuerpo)
            i += 1
            continue

        # Lista con viñetas
        if stripped.startswith('- ') or stripped.startswith('* '):
            cuerpo = stripped[2:]
            try:
                p = doc.add_paragraph(style='List Bullet')
            except KeyError:
                p = doc.add_paragraph()
            _agregar_runs_con_negrita(p, cuerpo)
            i += 1
            continue

        # Línea vacía
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Texto normal: respetar URLs activas + negrita inline
        if URL_RE.search(linea):
            p = _agregar_parrafo_con_links(doc, linea)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(14)
            _agregar_runs_con_negrita(p, linea)
        i += 1


def generar_docx_bytes(
    texto: str,
    config_org: Optional[dict] = None,
    tipo_doc: str = "",
    numero_doc: str = "",
) -> bytes:
    """
    Genera un .docx con membrete institucional.

    config_org: dict con nombre_organizacion, siglas, ciudad, anno_oficial.
    tipo_doc:   "oficio" / "carta" / "memorandum" / ... (ajusta márgenes).
    numero_doc: ej. "OFICIO N° 045-2026-SIGLAS".

    Si python-docx no está disponible, devuelve el texto plano en UTF-8.
    """
    texto = _limpiar_delimitadores(texto)

    if not _HAS_DOCX:
        return (texto or "").encode("utf-8")

    doc = Document()

    # Márgenes por tipo de documento
    cfg = MARGENES_POR_TIPO.get(
        (tipo_doc or "carta").lower(),
        {"izq": 3.0, "der": 2.5},
    )
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(cfg["izq"])
        section.right_margin  = Cm(cfg["der"])

    # Fuente por defecto
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ─── MEMBRETE ───
    org = config_org or {}
    nombre_org = (org.get("nombre_organizacion") or "").strip()
    siglas     = (org.get("siglas") or "").strip()
    ciudad     = (org.get("ciudad") or "").strip()
    anno       = (org.get("anno_oficial") or "").strip()

    if nombre_org:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(nombre_org)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1C, 0x3F, 0x8F)

    if siglas or ciudad:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = " — ".join([x for x in (siglas, ciudad) if x])
        r2 = p2.add_run(sub)
        r2.font.size = Pt(10)

    if anno:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(f'"{anno}"')
        run3.italic = True
        run3.font.size = Pt(9)

    # Línea separadora
    sep = doc.add_paragraph("─" * 80)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sep.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x1C, 0x3F, 0x8F)

    # Número de documento
    if numero_doc:
        p_num = doc.add_paragraph()
        p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_num = p_num.add_run(numero_doc)
        run_num.bold = True
        run_num.font.size = Pt(11)

    # ─── CUERPO ───
    doc.add_paragraph("")
    if es_contenido_html(texto):
        html_a_docx(doc, texto)
    else:
        markdown_a_docx(doc, texto or "")

    # ─── PIE ───
    if nombre_org or ciudad:
        section = doc.sections[0]
        footer = section.footer
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(" — ".join([x for x in (nombre_org, ciudad) if x]))
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
